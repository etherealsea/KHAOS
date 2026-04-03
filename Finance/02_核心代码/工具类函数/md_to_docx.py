import os
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

def parse_markdown_to_docx(md_path, docx_path):
    """
    一个简单的 Markdown 到 Docx 转换器，专门为毕设报告优化格式。
    """
    if not os.path.exists(md_path):
        print(f"Error: File not found {md_path}")
        return

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = u'Times New Roman'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    doc.styles['Normal'].font.size = Pt(12)
    
    current_level = 0
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # 1. Headers
        if line.startswith('# '):
            # 标题 1 (居中，加大)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[2:])
            run.bold = True
            run.font.size = Pt(22)
            run.font.name = u'黑体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
            
        elif line.startswith('## '):
            # 标题 2
            p = doc.add_paragraph()
            run = p.add_run(line[3:])
            run.bold = True
            run.font.size = Pt(16)
            run.font.name = u'黑体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
            
        elif line.startswith('### '):
            # 标题 3
            p = doc.add_paragraph()
            run = p.add_run(line[4:])
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = u'黑体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
            
        # 2. List Items
        elif line.startswith('* ') or line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            
        elif line.startswith('1. ') or (len(line)>2 and line[0].isdigit() and line[1]=='.'):
            # Find where the text starts
            dot_idx = line.find('.')
            p = doc.add_paragraph(line[dot_idx+1:].strip(), style='List Number')

        # 3. Horizontal Rule
        elif line.startswith('---'):
            doc.add_paragraph('_' * 40).alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 4. Bold Text parsing (Simple)
        else:
            p = doc.add_paragraph()
            # Split by ** for bolding
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    # Save
    doc.save(docx_path)
    print(f"Successfully converted {md_path} to {docx_path}")

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python md_to_docx.py <input_md> <output_docx>")
    else:
        parse_markdown_to_docx(sys.argv[1], sys.argv[2])
